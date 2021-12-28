from bs4 import BeautifulSoup
import pandas as pd
import os
import time
from simple_colors import *
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

time_start = time.time()
os.chdir(r"E:\OneDrive\mekansal_planlama\TUCBS\genisletme\xml\jeoteknik_final\prep")
file = open("jeoteknikEtud.v031.xml", mode="r")

def FindTurkishString(tempString):   
    turkishString = ["ş", "ö", "ü", "ç", "ı", "ğ", "İ", "Ş", "Ğ", "Ö" , "Ç", "Ü", "?"]    
    for i in turkishString:
        if i in tempString:
            #print("found turkish character:")
            return True

def FindIfFirstCharUpper (tempString):    
    if tempString[1].isupper()==True:
        #print("found first character upper")
        return True

def ConvertTurkishCharacter(tempString):
    
       
    try:
        tempString = tempString.replace("&#252", "ü").replace(';', '')
        tempString = tempString.replace("&#231", "ç").replace(';', '')
        tempString = tempString.replace("&#246", "ö").replace(';', '')
        tempString = tempString.replace("&#214", "Ö").replace(';', '')
        tempString = tempString.replace("&#199", "Ç").replace(';', '')
        tempString = tempString.replace('Ä±', 'ı')
        tempString = tempString.replace('ÅŸ', 'ş')
        tempString = tempString.replace('ÄŸ', 'ğ')
        tempString = tempString.replace('Ä°', 'İ')
        tempString = tempString.replace('Åž', 'Ö')
        tempString = tempString.replace('&#220', 'Ü')
        tempString = tempString.replace('&#226', 'â')
        
   
    except:
        None
   
    return tempString
    
content = file.read()
soup = BeautifulSoup(content, "lxml")
soup.encode("utf-8")

### Variables ###

codeList = list()
obj_list = list()
attributeList = list()
temp_str = str()
elementStereotypeCodeList=str()
elementStereotypeFeatureType=str()


dataType = dict()
featureTypeExp = dict()
variable_dict=dict()


dfTotal = pd.DataFrame()
dfFeatureTypeExpTotal = pd.DataFrame()
dfFeatureTypeExp = pd.DataFrame()

document = Document()


#list of codeList#

for element in soup.find_all("element", attrs={"xmi:type": "uml:Class"}):
    for a in element.find_all("properties", attrs={"stereotype": "CodeList"}):
        codeList.append(element.get("name"))


## Attributes ##
for element in soup.find_all("element", attrs={"xmi:type": "uml:Class"}):
    element_name = element.get("name")
    public = element.get("scope")
    variable_dict.update({"element_name" : element_name})
    variable_dict.update({"public" : public})
    
    element_name = ConvertTurkishCharacter(element_name)
  
    if FindTurkishString(element_name) == True:
        print(element_name)
  
    element_xmi_idref = element.get("xmi:idref")
    variable_dict.update({"element_xmi_idref" : element_xmi_idref})

    #print(element_name)
       
    elementStereotype = None
    elementStereotypeCodeList=str()
    elementStereotypeFeatureType=str()
    

    for b in element.find_all("properties", attrs={"stereotype": "FeatureType"}):
        elementStereotypeFeatureType = b.get("stereotype")
    
    for c in element.find_all("properties", attrs={"stereotype": "CodeList"}):
        elementStereotypeCodeList = c.get("stereotype")
            
    if elementStereotypeFeatureType == "":
        variable_dict.update({"elementStereotype" : elementStereotypeCodeList})
        elementStereotype = elementStereotypeCodeList
        #print("---" + elementStereotype)
    
    elif elementStereotypeCodeList == "":
        variable_dict.update({"elementStereotype" : elementStereotypeFeatureType})
        elementStereotype = elementStereotypeFeatureType
        #print("---" + elementStereotype)
  
    else:
        elementStereotype = None   
    
    element_explanation = str()
    
    for properities in element.find_all("properties", attrs={"stereotype": "FeatureType"}):

        element_explanation = properities.get("documentation")
        element_explanation = ConvertTurkishCharacter((element_explanation))
        variable_dict.update({"element_explanation " : element_explanation })
        
    for properities in element.find_all("properties", attrs={"stereotype": "CodeList"}):

        element_explanation = properities.get("documentation")
        element_explanation = ConvertTurkishCharacter((element_explanation))
        variable_dict.update({"element_explanation " : element_explanation })
        
   
    ### Creates Word Document Header
    document.add_heading(element.get("name"), level=5)
    table = document.add_table(rows =1, cols=1)
    table.style = 'TableGrid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = element.get("name")
    
    row_cells = table.add_row().cells
    row_cells[0].text = f"Ana paket: 	Jeoloji\n\
Tanım:\n\
{element_explanation}\n\
\n\
\
Stereotip:	<<{elementStereotype}>>"

  
    for j in element.find_all("attribute"):
        attribute_name = j.get("name")
        attPublic = j.get("scope")
               
        if FindTurkishString(repr(j.get("name"))) == True:
            print(green(element.get("name")))    
            print(j.get("name"))
        
        if FindIfFirstCharUpper(repr(j.get("name"))) == True:
            None
            #print(green(element.get("name")))   
            #print(j.get("name"))
        
        #Data type
        for c in j.find_all("properties"):
            attribute_type= c.get("type")
            if attribute_type == None:
                attribute_type =""
                
            attributeList.append(c.get("type"))
        
        #EA_Guid
        for k in repr(j).split("\n"):
            if "ea_guid" in k:
                ea_guid = k.split('"')[1].replace("}","").replace("{","")
                attribute_id = (k.split('"')[1].replace("}","").replace("{",""))
                 
        #Stereotype
        for strtype in j.find_all("stereotype"):
            #print(stereotype.get("stereotype"))
            stereotype = strtype.get("stereotype")
            if ((stereotype) == None):
                stereotype = ""
            #print(stereotype)
        
        
        for attrExp in j.find_all("documentation"):
            attributeExp = attrExp.get("value")
            attributeExp = ConvertTurkishCharacter(attributeExp)
            
        for inntValue in j.find_all("initial"):
            intValue = inntValue.get("body")
        
        for bound in j.find_all("bounds"):
            lowerBound = bound.get("lower")
            upperBound = bound.get("upper")
            multiplicity = f"[{lowerBound}..{upperBound}]"
            print(multiplicity)
            
            
            
           
#Word Table
        row_cells = table.add_row().cells
        row_cells[0].text = f"Öznitelik:	{attribute_name}\n\
Tipi:		{attribute_type}\n\
Tanım:\n\
{attributeExp}\n\
Çokluk: {multiplicity}\n\
Stereotip:{stereotype}\
"
              
        document.add_paragraph("")
          
        #variable_dict.update({"element_name" : element_name})
        #print("+++" + element_name)        
        variable_dict.update({"attribute_name" : attribute_name})
        variable_dict.update({"attribute_type": attribute_type})
        variable_dict.update({"attribute_id" : attribute_id})
        variable_dict.update({"stereotype" : stereotype})  
        variable_dict.update({"attributeExp" : attributeExp})  
        variable_dict.update({"intValue" : intValue})
        variable_dict.update({"attPublic" : attPublic})
  
        df = pd.DataFrame.from_dict(variable_dict, orient = "index")
        df = df.transpose()
        dfTotal = pd.concat([dfTotal, df],axis = 0)
        

dfTotal = dfTotal.reset_index()

try:
    dfTotal.to_excel("results.xlsx")
    document.save('demo.docx')
    
except:
    print("word is open")
    None
    
try:
    os.startfile('demo.docx')       
    os.startfile("results.xlsx")
    None  
     
except:
    print("excel is open")
    None

for i in attributeList:
    
    if i in codeList:
        codeList.remove(i)
    
time_end = time.time()
print(round((time_end - time_start),2))
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    